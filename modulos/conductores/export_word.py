import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def exportar(p, r, grafico_path=None):
    doc = Document()

    for _ in range(4):
        doc.add_paragraph()
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("MEMORIA DE CALCULO")
    run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Conductores Flexibles — IEEE 605 / IEC 60865-1")
    run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)

    doc.add_page_break()

    doc.add_heading("Parametros de Entrada", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, txt in enumerate(["Parametro", "Valor", "Unidad"]):
        hdr[i].text = txt
        for p_h in hdr[i].paragraphs:
            p_h.runs[0].bold = True

    input_map = [
        ("Subconductores por fase (n)", p["n"], "--"),
        ("Diametro subconductor (d)", p["d_mm"], "mm"),
        ("Separacion subconductores (as)", p["as_m"], "m"),
        ("Masa lineal (m's)", p["m_s"], "kg/m"),
        ("Area seccion (As)", p["As_mm2"], "mm2"),
        ("Longitud del vano (l)", p["l"], "m"),
        ("Distancia entre fases (a)", p["a"], "m"),
        ("Longitud aisladores (li)", p["li"], "m"),
        ("Flecha estatica (fes)", p["fes"], "m"),
        ("Velocidad viento (V)", p["V"], "m/s"),
        ("Corriente CC 3F (I''k3)", p["Ik3_kA"], "kA"),
        ("Duracion CC (Tk1)", p["Tk1"], "s"),
    ]
    for nom, val, uni in input_map:
        row = table.add_row().cells
        row[0].text = nom; row[1].text = str(val); row[2].text = uni

    doc.add_heading("Resultados", level=1)

    doc.add_heading("Cargas Estaticas (Viento)", level=2)
    t2 = doc.add_table(rows=1, cols=3)
    t2.style = 'Table Grid'
    for i, txt in enumerate(["Parametro", "Valor", "Unidad"]):
        t2.rows[0].cells[i].text = txt
    static_res = [
        ("Peso conductor (Fc)", r["Fc"], "N/m"),
        ("Peso hielo (FI)", r["FI"], "N/m"),
        ("Peso total vertical", r["w_tot"], "N/m"),
        ("Fuerza de viento (FW)", r["FW"], "N/m"),
        ("Tension estatica (Fst)", r["Fst"], "N"),
        ("Angulo balanceo x viento", r["d_viento"], "deg"),
        ("Desplazamiento horiz. x viento", r["bh_viento"], "m"),
    ]
    for nom, val, uni in static_res:
        row = t2.add_row().cells
        row[0].text = nom; row[1].text = str(val); row[2].text = uni

    doc.add_heading("Efectos de Cortocircuito", level=2)
    t3 = doc.add_table(rows=1, cols=4)
    t3.style = 'Table Grid'
    for i, txt in enumerate(["Parametro", "IEC 60865-1", "IEEE 605", "Unidad"]):
        t3.rows[0].cells[i].text = txt

    sc_map = [
        ("Tension CC (Ft)", r["Ft_iec"], r["Ft_ieee"], "N"),
        ("Drop-back (Ff)", r["Ff_iec"], r["Ff_ieee"], "N"),
        ("Angulo max balanceo", r["d_max_iec"], r["dm_ieee"], "deg"),
        ("Desplazamiento horiz. (bh)", r["bh_iec"], r["bh_ieee"], "m"),
        ("Distancia minima (amin)", r["a_min_iec"], r["a_min_ieee"], "m"),
    ]
    for nom, vi, vj, uni in sc_map:
        row = t3.add_row().cells
        row[0].text = nom; row[1].text = str(vi); row[2].text = str(vj); row[3].text = uni

    if grafico_path:
        doc.add_heading("Grafico 3D del Conductor", level=2)
        doc.add_picture(grafico_path, width=Inches(5.5))

    doc.add_paragraph()
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_ref.add_run("Referencias: IEEE Std 605-2023, IEC 60865-1:2011, IEC 60826:2017")
    run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
