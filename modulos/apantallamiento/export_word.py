import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def exportar(p, r, grafico_path=None):
    doc = Document()
    for _ in range(4):
        doc.add_paragraph()
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_tit.add_run("MEMORIA DE CALCULO")
    run.bold = True; run.font.size = Pt(24); run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sub.add_run("Apantallamiento (Esfera Rodante) — IEEE 998-2012")
    run.font.size = Pt(14); run.font.color.rgb = RGBColor(0x4A, 0x6F, 0xA5)
    doc.add_page_break()

    doc.add_heading("Parametros de Entrada", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    for i, txt in enumerate(["Parametro", "Valor", "Unidad"]):
        table.rows[0].cells[i].text = txt

    tipo_str = "Cable de guarda" if p.get("tipo_elemento", 1) == 1 else "Mastil / Punta Franklin"
    input_map = [
        ("Tension nominal (Vnom)", p.get("Vnom", 220), "kV"),
        ("BIL", p.get("BIL", 1050), "kV"),
        ("Altura conductor (h)", p.get("h", 15), "m"),
        ("Subconductores (n)", p.get("n", 2), "--"),
        ("Diametro subconductor (d)", p.get("d_mm", 25.4), "mm"),
        ("Separacion subconductores (s)", p.get("s_cm", 45.7), "cm"),
        ("Gradiente critico (E0)", p.get("E0", 21.1), "kV/cm"),
        ("Tipo elemento", tipo_str, ""),
    ]
    for nom, val, uni in input_map:
        row = table.add_row().cells
        row[0].text = nom; row[1].text = str(val); row[2].text = uni

    doc.add_heading("Resultados", level=1)
    t2 = doc.add_table(rows=1, cols=2)
    t2.style = 'Table Grid'
    for i, txt in enumerate(["Parametro", "Valor"]):
        t2.rows[0].cells[i].text = txt
    res_map = [
        ("Radio equivalente del haz (re)", f'{r.get("r_e","")} cm'),
        ("Radio corona Rc (Newton-Raphson)", f'{r.get("Rc","")} cm'),
        ("Iteraciones Newton-Raphson", str(r.get("n_iter_nr",""))),
    ]
    if r.get("Rc_prime"):
        res_map.append(("Radio corona corregido Rc'", f'{r["Rc_prime"]} cm'))
    res_map += [
        ("Impedancia de onda con corona (Zs)", f'{r.get("Zs","")} Ohm'),
        ("Corriente minima de blindaje (I_min)", f'{r.get("I_min","")} kA'),
        ("Radio de esfera rodante (S)", f'{r.get("S_esfera","")} m'),
    ]
    for nom, val in res_map:
        row = t2.add_row().cells
        row[0].text = nom; row[1].text = val

    if grafico_path:
        doc.add_heading("Grafico 3D", level=2)
        doc.add_picture(grafico_path, width=Inches(5.5))

    doc.add_paragraph()
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_ref.add_run("Referencia: IEEE Std 998-2012 — Guide for Direct Lightning Stroke Shielding of Substations")
    run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
